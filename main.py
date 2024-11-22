pip install requests python-docx
pip install requests beautifulsoup4 openai langchain-openai
pip install -U langchain-openai

import requests
from docx import Document
import os

subscription_key = "{key}"
endpoint = "https://api.cognitive.microsofttranslator.com/"
location = "eastus2"
target_language = "pt-br"

def translator_text(text, target_language):
  path = "/translate"
  constructed_url = endpoint + path
  headers = {
    'Ocp-Apim-Subscription-Key': subscription_key,
    'Ocp-Apim-Subscription-Region': location,
    'Content-type': 'application/json',
    'X-clientTraceId': str(os.urandom(16))
  }
  body=[{
          'text':text
  }]
  params= {
      'api-version':'3.0',
      'from':'en',
      'to':target_language
  }

  request = requests.post(constructed_url, params=params, headers=headers, json=body)
  response = request.json()
  return response[0]["translations"][0]["text"]

def translate_document(path, target_language):
    document = Document(path)
    full_text = []

    for paragraph in document.paragraphs:
      translated_text = translator_text(paragraph.text, target_language)
      full_text.append(translated_text)


    translated_doc = Document()
    for line in full_text:
        translated_doc.add_paragraph(line)

    path_translated = path.replace(".docx", f"_{target_language}.docx")
    translated_doc.save(path_translated)

    return path_translated

input_file = '/content/lyrics to translate.docx'
translate_document(input_file, target_language)

import requests
from bs4 import BeautifulSoup

def extract_text_from_url (url):
  response = requests.get(url)
  if response.status_code != 200:
    print(f"Failed to fetch the url. status code: {response.status_code}")
    return None

  soap = BeautifulSoup(response.text, 'html.parser')
  text = soap.get_text()
  for scipt_or_style in soap(['script', 'style']):
    scipt_or_style.decompose()
  texto = soap.get_text(separator=' ')
  lines = (line.strip() for line in texto.splitlines())
  chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
  text = '\n'.join(chunk for chunk in chunks if chunk)

  return text

extract_text_from_url('https://dev.to/westernal/a-letter-to-jobseekers-dont-give-up-2jim')


from langchain_openai import AzureChatOpenAI

client = AzureChatOpenAI(
  azure_endpoint="https://oai-ia102-translator-study.openai.azure.com/",
  api_key="{key}",
  api_version="2024-02-15-preview",
  deployment_name="gpt-4o-mini",
  max_retries=0
)

def translate_article(text, lang):
  messages = [
      ("system", "Você atua como tradutor de textos, por favor traduza o texto a seguir para o idioma especificado."),
      ("user", f"Traduza o seguinte texto: '{text}' para o idioma {lang} e responda em Markdown.")
  ]

  response = client.invoke(messages)
  return response.content

translate_article("Let's Begin with My Story", "português")

url = 'https://dev.to/westernal/a-letter-to-jobseekers-dont-give-up-2jim'
text = extract_text_from_url(url);
article = translate_article(text,'pt-br')